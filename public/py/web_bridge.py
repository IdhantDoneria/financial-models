"""Browser-side bridge between the terminal UI and the real model classes.

Runs inside Pyodide. The front end calls :func:`run_model` with a mnemonic and
a flat dict of slider values; this module translates those scalars into each
model's constructor arguments (synthesizing series where a model consumes
series — covariance assembly for MPT, seeded return draws for VaR, factor-
history windowing for Fama-French) and returns results, the ``explain()``
markdown and the ``visualize()`` figure as JSON.

The model classes themselves are imported *unchanged* from ``src/`` — this file
contains zero pricing/valuation logic.
"""

from __future__ import annotations

import json
import math
import time
from typing import Any, Callable

import numpy as np

from src import (
    BinomialTreeModel,
    BlackScholesModel,
    CAPMModel,
    DiscountedCashFlowModel,
    FamaFrenchModel,
    GordonGrowthModel,
    HestonModel,
    ModernPortfolioTheoryModel,
    MonteCarloOptionModel,
    ValueAtRiskModel,
)

_TRADING_DAYS = 252
_SEED = 20260702  # deterministic demo draws: same sliders -> same numbers


# --------------------------------------------------------------------------- #
# Builders: flat slider params -> model instance (+ optional extra outputs)
# --------------------------------------------------------------------------- #
def _build_dcf(p: dict) -> DiscountedCashFlowModel:
    # Project the FCF path from a base amount and constant growth ($M units;
    # shares are in millions, so per-share values land in plain dollars).
    years = int(p["years"])
    fcfs = [p["base_fcf"] * (1.0 + p["fcf_growth"]) ** t for t in range(1, years + 1)]
    return DiscountedCashFlowModel(
        free_cash_flows=fcfs,
        discount_rate=p["discount_rate"],
        terminal_growth=p["terminal_growth"],
        net_debt=p["net_debt"],
        shares_outstanding=p["shares_outstanding"],
    )


def _build_gg(p: dict) -> GordonGrowthModel:
    return GordonGrowthModel(
        dividend=p["dividend"],
        required_return=p["required_return"],
        growth=p["growth"],
    )


def _build_mpt(p: dict) -> ModernPortfolioTheoryModel:
    mu = np.array([p["mu1"], p["mu2"], p["mu3"]])
    sig = np.array([p["sigma1"], p["sigma2"], p["sigma3"]])
    rho = max(float(p["rho"]), -0.45)  # keep the 3-asset matrix positive-definite
    corr = np.full((3, 3), rho)
    np.fill_diagonal(corr, 1.0)
    cov = corr * np.outer(sig, sig)
    return ModernPortfolioTheoryModel(
        expected_returns=mu, covariance=cov, risk_free_rate=p["risk_free_rate"]
    )


def _build_var(p: dict) -> ValueAtRiskModel:
    # Scale annual mu/sigma to daily; historical & MC methods consume a seeded
    # synthetic daily return sample, parametric takes the moments directly.
    mu_d = p["mu_annual"] / _TRADING_DAYS
    sd_d = p["sigma_annual"] / math.sqrt(_TRADING_DAYS)
    common = dict(
        confidence_level=p["confidence"],
        horizon_days=int(p["horizon_days"]),
        portfolio_value=p["portfolio_value"],
        method=p["method"],
    )
    if p["method"] == "parametric":
        return ValueAtRiskModel(mean=mu_d, std=sd_d, **common)
    rng = np.random.default_rng(_SEED)
    returns = rng.normal(mu_d, sd_d, size=10 * _TRADING_DAYS)
    return ValueAtRiskModel(returns=returns, **common)


def _build_capm(p: dict) -> CAPMModel:
    return CAPMModel(
        risk_free_rate=p["risk_free_rate"],
        expected_market_return=p["expected_market_return"],
        beta=p["beta"],
    )


def _build_ff3(p: dict) -> tuple[FamaFrenchModel, dict]:
    # Real Ken French history (bundled snapshot) + user-chosen "true" loadings
    # -> synthetic asset returns; the regression must recover the loadings.
    factors = FamaFrenchModel.load_factors().tail(int(p["window"]))
    rng = np.random.default_rng(_SEED)
    eps = rng.normal(0.0, p["idio_sigma"], size=len(factors))
    r = (factors["RF"] + p["alpha"] + p["b_mkt"] * factors["Mkt-RF"]
         + p["s_smb"] * factors["SMB"] + p["h_hml"] * factors["HML"] + eps)
    extras = {
        "true_alpha": p["alpha"], "true_b_mkt": p["b_mkt"],
        "true_s_smb": p["s_smb"], "true_h_hml": p["h_hml"],
        "sample_start": str(factors.index[0]), "sample_end": str(factors.index[-1]),
    }
    return FamaFrenchModel(asset_returns=r.to_numpy(), factors=factors), extras


def _build_bsm(p: dict) -> BlackScholesModel:
    return BlackScholesModel(
        spot=p["spot"], strike=p["strike"], rate=p["rate"], sigma=p["sigma"],
        maturity=p["maturity"], option_type=p["option_type"],
        dividend_yield=p["dividend_yield"],
    )


def _build_crr(p: dict) -> BinomialTreeModel:
    return BinomialTreeModel(
        spot=p["spot"], strike=p["strike"], rate=p["rate"], sigma=p["sigma"],
        maturity=p["maturity"], option_type=p["option_type"],
        exercise=p["exercise"], dividend_yield=p["dividend_yield"],
        n_steps=int(p["n_steps"]),
    )


def _build_mc(p: dict) -> MonteCarloOptionModel:
    return MonteCarloOptionModel(
        spot=p["spot"], strike=p["strike"], rate=p["rate"], sigma=p["sigma"],
        maturity=p["maturity"], option_type=p["option_type"],
        dividend_yield=p["dividend_yield"], n_sims=int(p["n_sims"]),
        antithetic=bool(p["antithetic"]), seed=_SEED,
    )


def _build_hes(p: dict) -> HestonModel:
    return HestonModel(
        spot=p["spot"], strike=p["strike"], rate=p["rate"],
        maturity=p["maturity"], v0=p["v0"], kappa=p["kappa"],
        theta=p["theta"], xi=p["xi"], rho=p["rho"],
        option_type=p["option_type"],
    )


BUILDERS: dict[str, Callable[[dict], Any]] = {
    "DCF": _build_dcf, "GG": _build_gg, "MPT": _build_mpt, "VAR": _build_var,
    "CAPM": _build_capm, "FF3": _build_ff3, "BSM": _build_bsm,
    "CRR": _build_crr, "MC": _build_mc, "HES": _build_hes,
}


# --------------------------------------------------------------------------- #
# JSON sanitation: numpy scalars/arrays and non-finite floats -> JSON-safe
# --------------------------------------------------------------------------- #
def _clean(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(k): _clean(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [_clean(v) for v in value]
    if isinstance(value, np.ndarray):
        return [_clean(v) for v in value.tolist()]
    if isinstance(value, (np.floating, float)):
        f = float(value)
        return f if math.isfinite(f) else None
    if isinstance(value, (np.integer, int, bool)) or value is None:
        return value
    return str(value)


def run_model(mnemonic: str, params_json: str) -> str:
    """Build + run one model; returns a JSON payload for the front end."""
    params = json.loads(params_json)
    built = BUILDERS[mnemonic](params)
    model, extras = built if isinstance(built, tuple) else (built, {})

    t0 = time.perf_counter()
    results = model.calculate()
    calc_ms = (time.perf_counter() - t0) * 1000.0

    try:
        figure = model.visualize().to_json()
    except Exception as exc:  # chart failure must not take down the numbers
        figure = None
        extras = {**extras, "figure_error": str(exc)}

    return json.dumps({
        "ok": True,
        "results": _clean(results),
        "extras": _clean(extras),
        "explain": model.explain(),
        "figure": figure,
        "calc_ms": round(calc_ms, 2),
    })


# =========================================================================== #
# IB DESK — PDF analyzer (upload -> extract -> assume -> run -> export)
#
# Pipeline imports live inside the functions: the analyzer's pure-Python
# backends (pypdf, pdfminer.six, reportlab, openpyxl, python-docx) are
# micropip-installed on first use, after the main terminal has booted.
# =========================================================================== #
import base64
import re as _re

_ANALYZER: dict[str, Any] = {"data": None, "report": None, "period": None}

#: Fields the UI reports as FOUND/MISSING (order = display order).
_KEY_FIELDS = (
    "company_name", "ticker", "fiscal_year", "revenue", "free_cash_flows",
    "net_income", "total_debt", "cash_and_equivalents", "shares_outstanding",
    "current_price", "dividend_per_share", "beta", "revenue_growth",
    "operating_margin", "tax_rate",
)

_QUARTERLY_RE = _re.compile(
    r"(?i)\b(10-Q|quarterly report|three months ended|for the quarter ended|"
    r"third quarter|first quarter|second quarter|fourth quarter)\b")
_ANNUAL_RE = _re.compile(
    r"(?i)\b(10-K|annual report|fiscal year ended|for the year ended|"
    r"twelve months ended|full[- ]year)\b")


def _detect_period(text: str) -> str:
    """Classify the filing as annual or quarterly from its own language."""
    q = len(_QUARTERLY_RE.findall(text))
    a = len(_ANNUAL_RE.findall(text))
    return "quarterly" if q > a else "annual"


def _annualise_quarterly(data: Any) -> None:
    """Scale quarterly *flow* figures to annual run-rates, in place.

    Stocks (debt, cash, shares, price, beta) are point-in-time and unchanged;
    flows (revenue, net income, FCF, dividend) are multiplied by 4 and the
    quarter-over-quarter growth rate is compounded to an annual rate.
    """
    for field_name in ("revenue", "net_income"):
        value = getattr(data, field_name)
        if value is not None:
            setattr(data, field_name, value * 4.0)
    data.free_cash_flows = [f * 4.0 for f in data.free_cash_flows]
    if data.dividend_per_share is not None:
        data.dividend_per_share *= 4.0
    if data.revenue_growth is not None:
        data.revenue_growth = (1.0 + data.revenue_growth) ** 4 - 1.0


#: Extraction fields the user may set by hand (numeric; ``net_debt`` is a
#  derived property and ``free_cash_flows`` a synthesised series, so neither
#  is directly editable — adjust their inputs instead).
_OVERRIDABLE_FIELDS = (
    "revenue", "net_income", "total_debt", "cash_and_equivalents",
    "shares_outstanding", "current_price", "dividend_per_share", "beta",
    "revenue_growth", "operating_margin", "tax_rate",
)


def _assumed_preview(data: Any) -> dict:
    """The exact numbers the IB bot will use for each MISSING field.

    Mirrors :class:`AutoAssumer`'s defaults (and calls its own synthesiser for
    the FCF path) so the UI can show what "AUTO-ASSUMED" actually means —
    e.g. a company with no debt shows an assumed 0, not a hidden guess.
    """
    from src.pipeline import AutoAssumer

    auto = AutoAssumer()
    spot = data.current_price or 100.0
    preview: dict[str, Any] = {}
    if data.current_price is None:
        preview["current_price"] = 100.0            # normalised units
    if data.beta is None:
        preview["beta"] = auto.default_beta
    if data.dividend_per_share is None:
        preview["dividend_per_share"] = round(0.02 * spot, 4)
    if data.revenue is None:
        preview["revenue"] = 100.0                  # synth-FCF base, normalised
    if data.revenue_growth is None:
        preview["revenue_growth"] = 0.05
    if data.operating_margin is None:
        preview["operating_margin"] = 0.15
    if data.tax_rate is None:
        preview["tax_rate"] = auto.tax
    if data.total_debt is None:
        preview["total_debt"] = 0.0
    if data.cash_and_equivalents is None:
        preview["cash_and_equivalents"] = 0.0
    if data.net_debt is None:
        preview["net_debt"] = 0.0
    if data.shares_outstanding is None:
        preview["shares_outstanding"] = 1_000_000   # VaR notional proxy
    if data.net_income is None:
        preview["net_income"] = None                # not consumed by any model
    if not data.free_cash_flows:
        preview["free_cash_flows"] = [round(f, 2) for f in auto._synth_fcfs(data, 0.09)]
    return _clean(preview)


def override_field(key: str, value: Any = None) -> str:
    """Manually set (or reset to auto) one extracted field.

    ``value`` is a number, or ``None``/empty to hand the field back to the
    auto-assumer. Returns the refreshed fields + assumed preview so the UI
    can re-render, and invalidates any computed report (assumptions changed).
    """
    data = _ANALYZER["data"]
    if data is None:
        return json.dumps({"ok": False, "error": "No PDF analysed yet."})
    if key not in _OVERRIDABLE_FIELDS:
        return json.dumps({"ok": False, "error": f"Field {key!r} is not manually editable."})
    try:
        val = None if value in (None, "") else float(value)
    except (TypeError, ValueError):
        return json.dumps({"ok": False, "error": "Enter a number."})
    setattr(data, key, val)
    _ANALYZER["report"] = None            # previous report used old assumptions
    return json.dumps({"ok": True, "fields": _clean(data.to_dict()),
                       "assumed": _assumed_preview(data)})


def get_assumed() -> str:
    """Assumed-value preview for the current extraction (history restores)."""
    data = _ANALYZER["data"]
    if data is None:
        return json.dumps({"ok": False, "error": "No extraction loaded."})
    return json.dumps({"ok": True, "assumed": _assumed_preview(data)})


def analyze_pdf(pdf_bytes: Any, period_mode: str = "auto") -> str:
    """Extract financials from an uploaded PDF; returns a JSON payload.

    Args:
        pdf_bytes: The raw PDF (JS ``Uint8Array`` proxy or Python bytes).
        period_mode: ``"auto"`` (detect from the filing's language),
            ``"annual"`` or ``"quarterly"``.
    """
    from src.pipeline import PDFExtractor

    raw = bytes(pdf_bytes.to_py()) if hasattr(pdf_bytes, "to_py") else bytes(pdf_bytes)
    try:
        data = PDFExtractor().extract(raw)
    except Exception as exc:
        return json.dumps({"ok": False, "error": f"{type(exc).__name__}: {exc}"})

    period = period_mode if period_mode in ("annual", "quarterly") \
        else _detect_period(data.raw_text)
    if period == "quarterly":
        _annualise_quarterly(data)

    _ANALYZER.update(data=data, report=None, period=period)
    fields = _clean(data.to_dict())
    missing = [k for k in _KEY_FIELDS
               if fields.get(k) in (None, [], "") and k != "free_cash_flows"
               or (k == "free_cash_flows" and not fields.get(k))]
    return json.dumps({
        "ok": True, "fields": fields, "missing": missing, "period": period,
        "backends": fields.get("backends_used", []),
        "assumed": _assumed_preview(data),   # what AUTO-ASSUMED will really use
    })


def run_report(params_json: str) -> str:
    """Build assumptions (auto or manual) and run the selected models.

    ``params_json``: ``{"mode": "auto"|"manual", "selected": [names],
    "live_rf": float|null, "rf_source": str, "erp": float|null,
    "country": str, "currency": str, "overrides": {field: value}}``.

    ``live_rf`` and ``erp`` carry the SELECTED COUNTRY's live 10-year
    sovereign yield and Damodaran equity risk premium from the browser, so an
    Indian filing analysed under the India market uses India's cost of
    capital — never a hardcoded US base case.
    """
    from src.pipeline import (
        AnalysisRunner, AutoAssumer, ManualAssumer, ManualOverrides,
    )

    p = json.loads(params_json)
    data = _ANALYZER["data"]
    if data is None:
        return json.dumps({"ok": False, "error": "No PDF analysed yet — upload a filing first."})

    auto_kwargs = {}
    if p.get("live_rf") is not None:
        auto_kwargs["risk_free_rate"] = float(p["live_rf"])
    if p.get("erp") is not None:
        auto_kwargs["equity_risk_premium"] = float(p["erp"])
    auto = AutoAssumer(**auto_kwargs)

    if p.get("mode") == "manual":
        allowed = set(ManualOverrides.__dataclass_fields__)
        raw_overrides = {k: v for k, v in (p.get("overrides") or {}).items()
                         if k in allowed and v is not None}
        for int_field in ("var_horizon_days", "monte_carlo_paths"):
            if int_field in raw_overrides:
                raw_overrides[int_field] = int(raw_overrides[int_field])
        assumptions = ManualAssumer(auto).build(data, ManualOverrides(**raw_overrides))
        mode = "manual"
    else:
        assumptions = auto.build(data)
        mode = "auto"

    report = AnalysisRunner(data).run(assumptions, list(p.get("selected") or []), mode)
    _ANALYZER["report"] = report

    summary = report.summary_frame().to_dict(orient="records")
    rationale = {f"{model} · {param}": text
                 for (model, param), text in assumptions.rationale.items()}
    # Audit trail: which market the numbers were built in, in what currency.
    market_context = dict(assumptions.market_context)
    for extra_key in ("country", "currency", "fx_per_usd"):
        if p.get(extra_key) not in (None, ""):
            market_context[extra_key] = p[extra_key]
    if p.get("erp") is not None:
        market_context["equity_risk_premium"] = float(p["erp"])
    return json.dumps({
        "ok": True, "mode": mode, "summary": summary,
        "results": _clean(report.results), "errors": report.errors,
        "market_context": _clean(market_context),
        "rationale": rationale,
        "currency_symbol": p.get("currency_symbol") or "$",
        "rf_source": p.get("rf_source") or "default (Damodaran base case 4.25%)",
    })


def restore_extraction(fields_json: str, period: str = "annual") -> str:
    """Rehydrate a saved extraction (browser history) into the analyzer state.

    Saved analyses live in ``localStorage``; on reopen the UI shows the stored
    snapshot, but ``run_report``/``export_report`` need the Python-side
    ``ExtractedFinancials`` object back. ``fields_json`` is exactly what
    ``analyze_pdf`` returned (post-annualisation, so no re-scaling here).
    """
    from src.pipeline.pdf_extractor import ExtractedFinancials

    try:
        fields = json.loads(fields_json)
        allowed = set(ExtractedFinancials.__dataclass_fields__)
        kwargs = {k: v for k, v in fields.items() if k in allowed}
        if not kwargs.get("free_cash_flows"):
            kwargs["free_cash_flows"] = []
        data = ExtractedFinancials(**kwargs)
    except Exception as exc:
        return json.dumps({"ok": False, "error": f"{type(exc).__name__}: {exc}"})
    _ANALYZER.update(
        data=data, report=None,
        period=period if period in ("annual", "quarterly") else "annual")
    return json.dumps({"ok": True, "company": data.company_name})


def _fmt_docx(value: Any) -> str:
    """Human formatting for report values (mirrors the terminal grid)."""
    if isinstance(value, float):
        if abs(value) >= 1e5:
            return f"{value:,.0f}"
        if abs(value) >= 100:
            return f"{value:,.2f}"
        return f"{value:.6g}"
    if isinstance(value, list):
        return f"series · {len(value)} pts"
    return str(value)


def _build_docx(report: Any, path: str) -> None:
    """Write the analysis as a .docx — the format Google Docs imports natively."""
    import docx  # python-docx (lxml comes from the Pyodide distribution)

    doc = docx.Document()
    company = report.company.company_name or "Uploaded company"
    doc.add_heading(f"Financial Model Report — {company}", level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Mode: {report.mode.upper()}   ·   Period basis: "
        f"{(_ANALYZER['period'] or 'annual').upper()}   ·   "
        f"Generated by FINMODELS Terminal (in-browser Python)").italic = True

    doc.add_heading("Extracted financials", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in report.company.to_dict().items():
        if key == "backends_used" or value in (None, [], ""):
            continue
        cells = table.add_row().cells
        cells[0].text = key.replace("_", " ")
        cells[1].text = _fmt_docx(value)

    doc.add_heading("Assumptions (market context)", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in report.assumptions.market_context.items():
        cells = table.add_row().cells
        cells[0].text = key.replace("_", " ")
        cells[1].text = _fmt_docx(value)

    doc.add_heading("Model results", level=1)
    for model_name, results in report.results.items():
        doc.add_heading(model_name, level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for key, value in results.items():
            cells = table.add_row().cells
            cells[0].text = key.replace("_", " ")
            cells[1].text = _fmt_docx(value)

    if report.errors:
        doc.add_heading("Models not run", level=1)
        for model_name, err in report.errors.items():
            doc.add_paragraph(f"{model_name}: {err}")
    doc.save(path)


def export_report(fmt: str) -> str:
    """Render the last report as pdf / docx / xlsx; returns base64 JSON."""
    from src.pipeline import export_pdf, export_xlsx

    report = _ANALYZER["report"]
    if report is None:
        return json.dumps({"ok": False, "error": "Run a report before exporting."})

    company = (report.company.company_name or "company").strip()
    slug = _re.sub(r"[^A-Za-z0-9]+", "_", company).strip("_").lower() or "company"
    path = f"/tmp/{slug}_report.{fmt}"
    try:
        if fmt == "pdf":
            export_pdf(report, path)
            mime = "application/pdf"
        elif fmt == "xlsx":
            export_xlsx(report, path)
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif fmt == "docx":
            _build_docx(report, path)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            return json.dumps({"ok": False, "error": f"Unknown format {fmt!r}"})
    except Exception as exc:
        return json.dumps({"ok": False, "error": f"{type(exc).__name__}: {exc}"})

    with open(path, "rb") as fh:
        payload = base64.b64encode(fh.read()).decode()
    return json.dumps({"ok": True, "filename": f"{slug}_report.{fmt}",
                       "mime": mime, "b64": payload})
